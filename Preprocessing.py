import os
import docx


def Data_cleaning(text):
    from nltk.corpus import stopwords
    from nltk.tokenize import word_tokenize
    import string

    stop_words=set(stopwords.words('english'))

    words=word_tokenize(text)

    punctuated_sentence=[]

    for i in words:
        if i not in string.punctuation:
            punctuated_sentence.append(i)
        
    filtered_sentence=[]

    for i in punctuated_sentence:
        if i not in stop_words:
            filtered_sentence.append(i)

    from nltk.stem import PorterStemmer
    ps = PorterStemmer()
    stemmed_sentence=[]
    for w in filtered_sentence:
        stemmed_sentence.append(ps.stem(w))
    return stemmed_sentence

#READING ALL FILES...
text=''
file_count=0
path='G:\\project\\All Cancer  Files'
for filename in os.listdir(path):
    if filename.endswith('.docx'):
        print("Text len",len(text))
        if filename[:2] !='~$':
            file=path+"\\"+filename
            doc=docx.Document(file)
            i=0
            true=1
            file_count=file_count+1
            while true:
                try:
                    text=text+" "+doc.paragraphs[i].text
                    i=i+1
                except IndexError:
                    print(filename," Read Successfully...")
                    true=0
stemmed_sentence=Data_cleaning(text)
unique=list(set(stemmed_sentence))
row=file_count
column=len(unique)
from sklearn.feature_extraction.text import CountVectorizer

vectorizer = CountVectorizer()
print( vectorizer.fit_transform(text).todense() )
print( vectorizer.vocabulary_ )

#BOW BINARY VECTOR...

def BOW_Binary(text):
    temp=[0 for i in range(column)]
    col=0
    for word in unique:
        if word in pre_process:
            temp[col]=1
        else:
            temp[col]=0
        col=col+1
    return temp

#Trem Frequency
def Term_Frequency(text):
    wordfreq = []
    a=[]
    for w in unique:
        wordfreq.append(text.count(w))
    return wordfreq

#tf
def Tf(text):
    wordfreq = []
    a=[]
    for w in unique:
        wordfreq.append(text.count(w)/len(text))
    return wordfreq

#idf
def idf(Bow_Matrix):
    Idf_Matrix=[]
    import numpy as np
    import math
    a=list(np.shape(Bow_Matrix))
    for i in range(0,a[1]):
        temp=0
        for j in range(0,a[0]):
            temp=temp+Bow_Matrix[j][i]
        try:
            temp=math.log(a[0]/temp)
            Idf_Matrix.append(temp)
        except ZeroDivisionError:
            Idf_Matrix.append(0)
    return Idf_Matrix

#tf_idf
def tf_idf(Tf,idf):
    tfidf_Matrix=[]
    import numpy
    a=numpy.shape(Tf)
    for i in range(0,a[0]):
        temp_Matrix=[]
        for j in range(0,a[1]):
            temp_Matrix.append(Tf[i][j]*idf[j])
        tfidf_Matrix.append(temp_Matrix)
    return tfidf_Matrix

Bow_Matrix=[]
Term_Matrix=[]
Tf_Matrix=[]
Tf_Idf=[]
for filename in os.listdir(path):
    if filename.endswith('.docx'):
        if filename[:2] !='~$':
            text=''
            file=path+"\\"+filename
            doc=docx.Document(file)
            i=0
            true=1
            file_count=file_count+1
            while true:
                try:
                    text=text+doc.paragraphs[i].text
                    i=i+1
                except IndexError:
                    true=0
            pre_process=Data_cleaning(text)
            unique_pre_process=(list(set(pre_process)))
            Bow_Matrix.append(BOW_Binary(unique_pre_process))
            Term_Matrix.append(Term_Frequency(pre_process))
            Tf_Matrix.append(Tf(pre_process))
Idf_Matrix=idf(Bow_Matrix)
tfidf_matrix=tf_idf(Tf_Matrix,Idf_Matrix)
            
            
            
            
                    
                    

